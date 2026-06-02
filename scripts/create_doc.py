from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_para(text, bold=False, size=11, color=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def add_table_header(table, headers, bg='00466F'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        shade_cell(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)

# ══════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════
doc.add_paragraph()
title = doc.add_heading('NEURA Robotics - Meeting Automation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 70, 127)
    run.font.size = Pt(24)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Agentic Automation Prototype - Built with n8n + GPT-4o')
r.font.size = Pt(13); r.font.color.rgb = RGBColor(80, 80, 80); r.italic = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Prepared by Futurice   |   June 2026   |   PROTOTYPE / DEMO').font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════
# WHAT THIS IS
# ══════════════════════════════════════════════════
add_heading('What This Is', 1, color=(0, 70, 127))
add_para(
    'This document walks through a working prototype of an AI Meeting Agent built for the NEURA Robotics '
    'use case. The agent reads a meeting transcript, reasons over what was said, and produces structured '
    'outputs - action items, email drafts, a follow-up meeting invite, and a pattern check for recurring blockers.',
    size=11
)
doc.add_paragraph()
add_para(
    'The prototype was built in n8n using GPT-4o. It runs end to end today. '
    'A few things that require Microsoft 365 access are not yet connected - '
    'those are clearly marked below.',
    size=11
)
doc.add_paragraph()

t = doc.add_table(rows=2, cols=2)
t.style = 'Table Grid'
shade_cell(t.rows[0].cells[0], 'E8F5E9')
shade_cell(t.rows[0].cells[1], 'FFF8E1')
shade_cell(t.rows[1].cells[0], 'E8F5E9')
shade_cell(t.rows[1].cells[1], 'FFF8E1')

t.rows[0].cells[0].text = '✅  Working in this prototype'
t.rows[0].cells[1].text = '⏳  Needs M365 access to complete'
t.rows[1].cells[0].text = (
    '• AI reads and reasons over the meeting transcript\n'
    '• Extracts action items with owners, deadlines and priorities\n'
    '• Detects if the same blocker has appeared in past meetings\n'
    '• Detects external (non-NEURA) participants\n'
    '• Human approval step - fully live\n'
    '• Drafted emails ready to send\n'
    '• Follow-up meeting invite drafted'
)
t.rows[1].cells[1].text = (
    '• Actually sending the emails (drafted but not delivered)\n'
    '• Creating the calendar event in Teams (details generated but not created)\n'
    '• Reading from a live Teams transcript (a dummy transcript is used here)'
)

for row in t.rows:
    for i, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                if row == t.rows[0]:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 100, 0) if i == 0 else RGBColor(180, 100, 0)

doc.add_page_break()

# ══════════════════════════════════════════════════
# WORKFLOW CANVAS
# ══════════════════════════════════════════════════
add_heading('The Workflow', 1, color=(0, 70, 127))
add_para(
    'The screenshot below shows the complete workflow running in n8n. '
    'Each node with a green checkmark ran successfully. '
    'The agent node in the centre calls three tools autonomously, '
    'then routes through a human approval step on the right.',
    size=11
)
doc.add_paragraph()

screenshot_path = '/Users/ancy.joseph/Documents/Test/Teams meeting/workflow_screenshot.png'
if os.path.exists(screenshot_path):
    doc.add_picture(screenshot_path, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('[INSERT SCREENSHOT OF n8n CANVAS HERE]')
    r.bold = True; r.font.color.rgb = RGBColor(180, 0, 0); r.font.size = Pt(11)

doc.add_paragraph()
cap = doc.add_paragraph('Figure 1 - Meeting Automation workflow in n8n (all nodes ran successfully)')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in cap.runs:
    run.font.size = Pt(9); run.italic = True; run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ══════════════════════════════════════════════════
# HOW IT WORKS
# ══════════════════════════════════════════════════
add_heading('How It Works', 1, color=(0, 70, 127))
add_para('The agent follows these steps every time it processes a meeting:', size=11)
doc.add_paragraph()

steps = [
    ('1. Reads the transcript',
     'The agent receives the full meeting transcript along with the attendee list. '
     'It reads every line and decides what actions to take - no templates, no rules, just reasoning over the conversation.'),
    ('2. Checks for recurring blockers',
     'For each problem or delay mentioned, the agent queries a history database to see if the same issue has come up before. '
     'In this demo it found that the Bosch API documentation delay had appeared multiple times in previous months and flagged it as a recurring risk.'),
    ('3. Extracts every action item',
     'The agent identifies every commitment made in the meeting - including implied ones. '
     'For example, Felix said "if I have the docs by June 6th, integration testing is done by June 14th." '
     'The agent captured this as an action item with a dependency on Thomas delivering the docs first.'),
    ('4. Drafts personalised emails',
     'For each action item owner, the agent drafts a personalised email with their specific task, deadline and priority. '
     'These are ready to send - the content is correct and complete. '
     'They are not sent yet because Microsoft 365 access is pending.'),
    ('5. Drafts a follow-up meeting invite',
     'The agent detected that everyone agreed to a follow-up review on June 13th and drafted a calendar invite '
     'with all attendees, the agenda and a placeholder Teams link. '
     'Not created in the calendar yet - pending M365 access.'),
    ('6. Detects external participants and pauses',
     'Thomas Becker is from Bosch - an external company. The agent detected this automatically '
     'and paused the entire workflow. Nothing was sent or distributed. '
     'It waited for a human to review and approve before proceeding.'),
    ('7. Human approves - workflow resumes',
     'The organiser opened the approval link, reviewed the summary and action items, and clicked Approve. '
     'The workflow resumed and completed.'),
]

for title_text, body_text in steps:
    add_para(title_text, bold=True, size=11, color=(0, 70, 127))
    add_para(body_text, size=11)
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════
# INPUT
# ══════════════════════════════════════════════════
add_heading('Input - Meeting Transcript', 1, color=(0, 70, 127))
add_para(
    'The transcript below is fictional, written specifically for this demo. '
    'In production this comes directly from the Teams meeting recording via Microsoft Graph API.',
    size=11, italic=True
)
doc.add_paragraph()
add_para('Meeting:  Robot Arm Integration Sync - Sprint Review  |  1 June 2026', bold=True)
doc.add_paragraph()
add_para('Attendees:', bold=True)
add_bullet('Maria Schmidt  -  maria.schmidt@neura-robotics.com  [Internal, Organiser]')
add_bullet('Felix Wagner  -  felix.wagner@neura-robotics.com  [Internal]')
add_bullet('Jonas Müller  -  jonas.mueller@neura-robotics.com  [Internal]')
add_bullet('Thomas Becker  -  thomas.becker@bosch.com  [External - Bosch]')
doc.add_paragraph()
add_para('Transcript:', bold=True)

transcript_lines = [
    ('Maria Schmidt', 'Good morning everyone. Sprint review, then integration issues. Felix - firmware status?'),
    ('Felix Wagner',  '85% done. Main blocker is the same IMU sensor calibration issue from last sprint. The drift is affecting precision targeting by 2.3 millimetres. I need the updated calibration constants from Jonas by Thursday.'),
    ('Jonas Müller',  'I can get them to you Wednesday. I have the root cause - temperature coefficient issue in the MEMS sensor. I will send the constants and a test protocol.'),
    ('Maria Schmidt', 'Good. Thomas - your team was supposed to deliver the integration API documentation two weeks ago. We have now been blocked twice waiting on this. When can we actually commit?'),
    ('Thomas Becker', 'I understand the frustration. We had internal re-scoping at Bosch. I personally commit to delivering full API documentation by June 6th. I will also run a technical walkthrough call that same week.'),
    ('Maria Schmidt', 'That is the third time the API docs have slipped, Thomas. Felix - once you have the docs, how long to complete integration?'),
    ('Felix Wagner',  'If I have the docs by June 6th, integration testing is done by June 14th. But I also need the hardware interface specs, not just the software API.'),
    ('Thomas Becker', 'Hardware specs are ready now. I will send them today before end of day.'),
    ('Maria Schmidt', 'Jonas - can you also investigate the power consumption regression in build 2.4? The arm is drawing 340W at peak versus our 280W target.'),
    ('Jonas Müller',  'Yes. I suspect the servo controller update in the last build. I will have initial findings by Friday.'),
    ('Maria Schmidt', 'Let us schedule a follow-up integration review for June 13th - one week after Thomas delivers. Thomas must be on that call. Agreed?'),
    ('All',           'Agreed.'),
    ('Maria Schmidt', 'Good. Felix and Jonas - I am counting on you. Thomas - no more slippage. That is all.'),
]

for speaker, line in transcript_lines:
    p = doc.add_paragraph()
    r1 = p.add_run(speaker + ': ')
    r1.bold = True; r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0, 70, 127)
    p.add_run(line).font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════════════
# OUTPUT - SUMMARY & ACTION ITEMS
# ══════════════════════════════════════════════════
add_heading('Output - What the Agent Produced', 1, color=(0, 70, 127))
add_para('The following is the real structured output from the agent after processing the transcript and receiving human approval.', size=11)
doc.add_paragraph()

add_para('Status', bold=True, color=(0, 70, 127))
add_para('COMPLETED - Human Approved', size=11)
doc.add_paragraph()

add_para('Meeting Summary', bold=True, color=(0, 70, 127))
add_para(
    'The meeting focused on reviewing the sprint progress and addressing integration issues with the robot arm. '
    'Key blockers included the IMU sensor calibration and delayed API documentation from Bosch. '
    'A follow-up meeting was scheduled to ensure progress on these issues.',
    size=11
)
doc.add_paragraph()

add_para('Action Items', bold=True, color=(0, 70, 127))
doc.add_paragraph()

t_ai = doc.add_table(rows=1, cols=5)
t_ai.style = 'Table Grid'
t_ai.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(t_ai, ['Owner', 'Email', 'Task', 'Deadline', 'Priority'])

action_items = [
    ('Felix Wagner', 'felix.wagner@neura-robotics.com',
     'Complete firmware integration testing once API documentation is received',
     '2026-06-14', 'HIGH', 'F8F9FA', 'Thomas Becker delivers API docs by 2026-06-06'),
    ('Jonas Müller', 'jonas.mueller@neura-robotics.com',
     'Deliver updated calibration constants and test protocol to Felix',
     '2026-06-03', 'CRITICAL', 'FFF3CD', None),
    ('Thomas Becker\n(Bosch - External)', 'thomas.becker@bosch.com',
     'Deliver full API documentation and run technical walkthrough',
     '2026-06-06', 'CRITICAL', 'FFF3CD', None),
    ('Jonas Müller', 'jonas.mueller@neura-robotics.com',
     'Investigate power consumption regression in build 2.4',
     '2026-06-05', 'HIGH', 'F8F9FA', None),
]

for owner, email, task, deadline, priority, bg, depends in action_items:
    row = t_ai.add_row()
    row.cells[0].text = owner
    row.cells[1].text = email
    row.cells[2].text = task + (('\n🔗 Depends on: ' + depends) if depends else '')
    row.cells[3].text = deadline
    row.cells[4].text = priority
    for cell in row.cells:
        shade_cell(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                if cell == row.cells[4] and priority == 'CRITICAL':
                    run.font.color.rgb = RGBColor(180, 0, 0); run.bold = True

doc.add_paragraph()

add_para('Follow-up Meeting (drafted - not yet created in calendar)', bold=True, color=(0, 70, 127))
fm = doc.add_table(rows=5, cols=2)
fm.style = 'Table Grid'
for i, (label, value) in enumerate([
    ('Title',     'Integration Review Follow-up'),
    ('Date',      'Friday 13 June 2026 at 09:00 (60 minutes)'),
    ('Attendees', 'maria.schmidt@neura-robotics.com\nfelix.wagner@neura-robotics.com\njonas.mueller@neura-robotics.com\nthomas.becker@bosch.com'),
    ('Agenda',    'Review integration progress and address any remaining issues'),
    ('Teams Link','[Generated by Microsoft Graph API in production]'),
]):
    fm.rows[i].cells[0].text = label
    fm.rows[i].cells[1].text = value
    shade_cell(fm.rows[i].cells[0], 'E8F0F7')
    for para in fm.rows[i].cells[0].paragraphs:
        for run in para.runs: run.bold = True; run.font.size = Pt(10)
    for para in fm.rows[i].cells[1].paragraphs:
        for run in para.runs: run.font.size = Pt(10)

doc.add_paragraph()
add_para('Risk Flag Detected', bold=True, color=(0, 70, 127))
add_bullet('Recurring delays in API documentation delivery from Bosch - the agent cross-referenced meeting history and identified this as a pattern, not a one-off.')

doc.add_page_break()

# ══════════════════════════════════════════════════
# EMAIL DRAFTS
# ══════════════════════════════════════════════════
add_heading('Drafted Emails', 1, color=(0, 70, 127))
add_para(
    'The agent drafted the following emails automatically. They are ready to send. '
    'In production, they would be delivered via Microsoft Outlook (Microsoft Graph API). '
    'They are not sent in this prototype because M365 access is pending.',
    size=11
)
doc.add_paragraph()

emails = [
    {
        'n': '1 of 5', 'type': 'Action Item Email',
        'to': 'felix.wagner@neura-robotics.com',
        'subject': '[Action Required] Complete firmware integration testing once API documentation is received - due 2026-06-14',
        'body': (
            'Hi Felix Wagner,\n\n'
            'Following today\'s Robot Arm Integration Sync - Sprint Review, you have been assigned the following action item:\n\n'
            '📌  Task:      Complete firmware integration testing once API documentation is received\n'
            '⏰  Deadline:  2026-06-14\n'
            '🎯  Priority:  HIGH\n'
            '🔗  Depends on: Thomas Becker delivers API documentation by 2026-06-06\n\n'
            'A follow-up review has been scheduled for 2026-06-13 at 09:00 to track progress.\n\n'
            'Please confirm receipt and flag any blockers as soon as possible.\n\n'
            '- NEURA AI Meeting Agent\n'
            '[DEMO - not sent. In production: Microsoft Graph API POST /me/sendMail]'
        )
    },
    {
        'n': '2 of 5', 'type': 'Action Item Email',
        'to': 'jonas.mueller@neura-robotics.com',
        'subject': '[Action Required] Deliver updated calibration constants and test protocol to Felix - due 2026-06-03',
        'body': (
            'Hi Jonas Müller,\n\n'
            'Following today\'s Robot Arm Integration Sync - Sprint Review, you have been assigned the following action item:\n\n'
            '📌  Task:      Deliver updated calibration constants and test protocol to Felix\n'
            '⏰  Deadline:  2026-06-03\n'
            '🎯  Priority:  CRITICAL\n\n'
            'A follow-up review has been scheduled for 2026-06-13 at 09:00 to track progress.\n\n'
            'Please confirm receipt and flag any blockers as soon as possible.\n\n'
            '- NEURA AI Meeting Agent\n'
            '[DEMO - not sent. In production: Microsoft Graph API POST /me/sendMail]'
        )
    },
    {
        'n': '3 of 5', 'type': 'Action Item Email',
        'to': 'thomas.becker@bosch.com',
        'subject': '[Action Required] Deliver full API documentation and run technical walkthrough - due 2026-06-06',
        'body': (
            'Hi Thomas Becker,\n\n'
            'Following today\'s Robot Arm Integration Sync - Sprint Review, you have been assigned the following action item:\n\n'
            '📌  Task:      Deliver full API documentation and run technical walkthrough\n'
            '⏰  Deadline:  2026-06-06\n'
            '🎯  Priority:  CRITICAL\n\n'
            'A follow-up review has been scheduled for 2026-06-13 at 09:00 to track progress.\n\n'
            'Please confirm receipt and flag any blockers as soon as possible.\n\n'
            '- NEURA AI Meeting Agent\n'
            '[DEMO - not sent. In production: Microsoft Graph API POST /me/sendMail]'
        )
    },
    {
        'n': '4 of 5', 'type': 'Action Item Email',
        'to': 'jonas.mueller@neura-robotics.com',
        'subject': '[Action Required] Investigate power consumption regression in build 2.4 - due 2026-06-05',
        'body': (
            'Hi Jonas Müller,\n\n'
            'Following today\'s Robot Arm Integration Sync - Sprint Review, you have been assigned the following action item:\n\n'
            '📌  Task:      Investigate power consumption regression in build 2.4\n'
            '⏰  Deadline:  2026-06-05\n'
            '🎯  Priority:  HIGH\n\n'
            'A follow-up review has been scheduled for 2026-06-13 at 09:00 to track progress.\n\n'
            'Please confirm receipt and flag any blockers as soon as possible.\n\n'
            '- NEURA AI Meeting Agent\n'
            '[DEMO - not sent. In production: Microsoft Graph API POST /me/sendMail]'
        )
    },
    {
        'n': '5 of 5', 'type': 'Follow-up Meeting Invite',
        'to': 'maria.schmidt@neura-robotics.com, felix.wagner@neura-robotics.com,\njonas.mueller@neura-robotics.com, thomas.becker@bosch.com',
        'subject': 'Meeting Invite: Integration Review Follow-up - 2026-06-13',
        'body': (
            'Hi all,\n\n'
            'A follow-up meeting has been scheduled. Please add it to your calendar.\n\n'
            '📅  Title:     Integration Review Follow-up\n'
            '🗓   Date:      2026-06-13 at 09:00\n'
            '⏱   Duration:  60 minutes\n'
            '📋  Agenda:    Review integration progress and address any remaining issues\n\n'
            'Attendees:\n'
            '  • maria.schmidt@neura-robotics.com\n'
            '  • felix.wagner@neura-robotics.com\n'
            '  • jonas.mueller@neura-robotics.com\n'
            '  • thomas.becker@bosch.com\n\n'
            '🔗  Teams Link: [Generated by Graph API in production]\n\n'
            '- NEURA AI Meeting Agent\n'
            '[DEMO - not sent. In production: Microsoft Graph API POST /me/calendar/events]'
        )
    },
]

for email in emails:
    add_para(f"Email {email['n']} - {email['type']}", bold=True, size=11, color=(0, 70, 127))
    et = doc.add_table(rows=3, cols=2)
    et.style = 'Table Grid'
    for i, (label, value) in enumerate([
        ('To', email['to']),
        ('Subject', email['subject']),
        ('Body', email['body']),
    ]):
        et.rows[i].cells[0].text = label
        et.rows[i].cells[1].text = value
        shade_cell(et.rows[i].cells[0], 'E8F0F7')
        for para in et.rows[i].cells[0].paragraphs:
            for run in para.runs: run.bold = True; run.font.size = Pt(9)
        for para in et.rows[i].cells[1].paragraphs:
            for run in para.runs: run.font.size = Pt(9)
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════
# HUMAN-IN-THE-LOOP
# ══════════════════════════════════════════════════
add_heading('Human Approval Step', 1, color=(0, 70, 127))
add_para(
    'One of the core features is the ability to pause the agent before it takes any action involving external people. '
    'This is not a fixed rule - it is a configurable condition that can be set to whatever makes sense for NEURA.',
    size=11
)
doc.add_paragraph()

add_para('How it worked in this demo:', bold=True)
add_bullet('The agent detected that Thomas Becker (Bosch) was an external participant')
add_bullet('The entire workflow paused - nothing was sent or distributed')
add_bullet('An approval link was generated and sent to the meeting organiser')
add_bullet('The organiser reviewed the summary and action items, then clicked Approve')
add_bullet('The workflow resumed and completed all actions')
add_bullet('If the organiser had clicked Reject, the minutes would have been saved as a draft with nothing sent')

doc.add_paragraph()
add_para('The approval trigger is fully configurable:', bold=True)
doc.add_paragraph()

t_hitl = doc.add_table(rows=1, cols=2)
t_hitl.style = 'Table Grid'
add_table_header(t_hitl, ['Trigger condition', 'When you would use this'])

trigger_rows = [
    ('External participant in the meeting', '✅  Current setting - matches the use case spec'),
    ('All meetings, always',                'Maximum control - nothing leaves without sign-off'),
    ('Specific person attended',            'E.g. a CEO, legal contact, or key client'),
    ('Agent detects sensitive content',     'E.g. budget figures, legal matters, confidential topics'),
    ('Meeting above a certain size',        'E.g. more than 10 people attended'),
    ('Any combination of the above',        'Mix and match conditions as needed'),
]

for condition, use_case in trigger_rows:
    row = t_hitl.add_row()
    row.cells[0].text = condition
    row.cells[1].text = use_case
    bg = 'E8F5E9' if '✅' in use_case else 'F8F9FA'
    shade_cell(row.cells[0], bg)
    shade_cell(row.cells[1], bg)
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs: run.font.size = Pt(9)

doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════
# CURRENT vs M365
# ══════════════════════════════════════════════════
add_heading('Current State vs Full Integration', 1, color=(0, 70, 127))
add_para(
    'The agent intelligence, approval flow and email drafting are all working today. '
    'The three items below require Microsoft 365 sandbox access, which is currently blocked '
    'by an IT policy. Once unblocked, connecting them is roughly one hour of work.',
    size=11
)
doc.add_paragraph()

t_state = doc.add_table(rows=1, cols=3)
t_state.style = 'Table Grid'
add_table_header(t_state, ['Capability', 'Today', 'With M365 access'])

state_rows = [
    ('Meeting transcript',          '⏳ Fictional dummy transcript',       'Live Teams recording via Graph API'),
    ('Action item extraction',      '✅ Fully working',                    '✅ No change'),
    ('Recurring blocker detection', '✅ Fully working',                    'Queries real meeting history database'),
    ('Email drafting',              '✅ Fully working - content is ready', '✅ No change - emails would be sent'),
    ('Sending the emails',          '⏳ Not sent - M365 pending',          'Real Outlook emails via Graph API'),
    ('Follow-up meeting invite',    '⏳ Drafted - not created in calendar', 'Real Teams calendar event via Graph API'),
    ('External participant check',  '✅ Fully working',                    'From real Teams attendance report'),
    ('Human approval flow',         '✅ Fully working',                    '✅ No change'),
]

for cap, today, m365 in state_rows:
    row = t_state.add_row()
    row.cells[0].text = cap
    row.cells[1].text = today
    row.cells[2].text = m365
    shade_cell(row.cells[1], 'FFF8E1' if '⏳' in today else 'E8F5E9')
    shade_cell(row.cells[2], 'E8F5E9')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs: run.font.size = Pt(9)

doc.add_paragraph()
add_para('Estimated time to connect M365 once access is available: approximately 1 hour.', bold=True, size=11, color=(0, 120, 0))

doc.add_page_break()

# ══════════════════════════════════════════════════
# TECH STACK
# ══════════════════════════════════════════════════
add_heading('Technology', 1, color=(0, 70, 127))

t_tech = doc.add_table(rows=6, cols=2)
t_tech.style = 'Table Grid'
for i, (label, value) in enumerate([
    ('Workflow engine',      'n8n (self-hosted, open source)'),
    ('AI model',             'GPT-4o via OpenRouter'),
    ('Agent type',           'Tools Agent - the AI decides which tools to call and when'),
    ('Tools',                'Check Blocker History / Schedule Follow-up Meeting / Send Action Email'),
    ('Human approval',       'n8n Wait node - pauses the workflow, resumes when a human clicks Approve'),
    ('Target integration',   'Microsoft Graph API - Mail.Send, Calendars.ReadWrite (M365 access pending)'),
]):
    t_tech.rows[i].cells[0].text = label
    t_tech.rows[i].cells[1].text = value
    shade_cell(t_tech.rows[i].cells[0], 'E8F0F7')
    for para in t_tech.rows[i].cells[0].paragraphs:
        for run in para.runs: run.bold = True; run.font.size = Pt(10)
    for para in t_tech.rows[i].cells[1].paragraphs:
        for run in para.runs: run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph()

final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = final.add_run(
    'This is a prototype built to demonstrate agentic automation capability in response to the NEURA Robotics RFI. '
    'All meeting content, names and email addresses are fictional and used for demonstration purposes only.'
)
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = RGBColor(120, 120, 120)

output_path = '/Users/ancy.joseph/Documents/Test/Teams meeting/NEURA_Meeting_Agent_Prototype.docx'
doc.save(output_path)
print('Saved:', output_path)
